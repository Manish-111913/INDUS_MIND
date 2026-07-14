"""Type-aware parsing → normalized units (docs/02 §10 step 2).

Produces `ParsedUnit`s (text blocks + atomic tables, each with page_no /
section_path / bbox) that the chunker consumes. PDFs use the OCR/text-layer
pages plus pdfplumber tables; xlsx via openpyxl, docx via python-docx, .msg via
extract-msg (attachments recursed). All heavy libs are lazy-imported.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field

from app.core.logging import get_logger
from app.core.ocr import OCRPage

log = get_logger("ingestion.parsing")

_HEADING = re.compile(r"^(?:\d+(?:\.\d+)*\s+.+|[A-Z][A-Z0-9 \-/]{4,})$")


@dataclass(slots=True)
class ParsedUnit:
    page_no: int | None
    section_path: str | None
    text: str
    kind: str = "text"  # text | table
    bbox: dict | None = None


@dataclass(slots=True)
class ParsedDocument:
    units: list[ParsedUnit] = field(default_factory=list)
    page_count: int = 0


def parse_document(mime: str, data: bytes, ocr_pages: list[OCRPage]) -> ParsedDocument:
    if mime == "application/pdf":
        return _parse_pdf(data, ocr_pages)
    if mime in {
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    }:
        return _parse_xlsx(data)
    if mime in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }:
        return _parse_docx(data)
    if mime in {"application/vnd.ms-outlook", "message/rfc822"}:
        return _parse_email(data, mime)
    if mime in {"text/plain", "text/csv"} or mime.startswith("image/"):
        return _parse_text(ocr_pages, data)
    log.warning("parse_unsupported_mime", mime=mime)
    return _parse_text(ocr_pages, data)


def _section_for(text: str, current: str | None) -> str | None:
    first = text.strip().splitlines()[0] if text.strip() else ""
    return first[:200] if _HEADING.match(first) else current


def _parse_pdf(data: bytes, ocr_pages: list[OCRPage]) -> ParsedDocument:
    units: list[ParsedUnit] = []
    section: str | None = None
    for page in ocr_pages:
        blocks = page.blocks or [_as_block(page.text)]
        for block in blocks:
            if not block.text.strip():
                continue
            section = _section_for(block.text, section)
            units.append(ParsedUnit(page_no=page.page_no, section_path=section,
                                    text=block.text.strip(),
                                    bbox={"coords": block.bbox} if getattr(block, "bbox", None) else None))
    units.extend(_pdf_tables(data))
    return ParsedDocument(units=units, page_count=len(ocr_pages))


def _pdf_tables(data: bytes) -> list[ParsedUnit]:
    try:
        import io

        import pdfplumber  # lazy
    except ImportError:
        return []
    units: list[ParsedUnit] = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for pno, page in enumerate(pdf.pages, start=1):
                for table in page.extract_tables() or []:
                    rows = ["\t".join((c or "") for c in row) for row in table if any(row)]
                    if rows:
                        units.append(ParsedUnit(page_no=pno, section_path="table",
                                                text="\n".join(rows), kind="table"))
    except Exception as exc:  # noqa: BLE001
        log.warning("pdf_table_extract_failed", error=str(exc))
    return units


def _parse_xlsx(data: bytes) -> ParsedDocument:
    import io

    from openpyxl import load_workbook  # lazy

    units: list[ParsedUnit] = []
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    for sheet in wb.worksheets:
        rows = ["\t".join("" if c is None else str(c) for c in row)
                for row in sheet.iter_rows(values_only=True)]
        rows = [r for r in rows if r.strip()]
        if rows:
            units.append(ParsedUnit(page_no=None, section_path=sheet.title,
                                    text="\n".join(rows), kind="table"))
    return ParsedDocument(units=units, page_count=len(wb.worksheets))


def _parse_docx(data: bytes) -> ParsedDocument:
    import io

    from docx import Document as Docx  # lazy (python-docx)

    doc = Docx(io.BytesIO(data))
    units: list[ParsedUnit] = []
    section: str | None = None
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        section = _section_for(text, section)
        units.append(ParsedUnit(page_no=None, section_path=section, text=text))
    for table in doc.tables:
        rows = ["\t".join(cell.text for cell in row.cells) for row in table.rows]
        if rows:
            units.append(ParsedUnit(page_no=None, section_path="table",
                                    text="\n".join(rows), kind="table"))
    return ParsedDocument(units=units, page_count=1)


def _parse_email(data: bytes, mime: str) -> ParsedDocument:
    units: list[ParsedUnit] = []
    try:
        import tempfile

        import extract_msg  # lazy
        with tempfile.NamedTemporaryFile(suffix=".msg", delete=False) as tmp:
            tmp.write(data)
            path = tmp.name
        msg = extract_msg.Message(path)
        header = f"From: {msg.sender}\nTo: {msg.to}\nSubject: {msg.subject}"
        units.append(ParsedUnit(page_no=None, section_path="email", text=header))
        units.append(ParsedUnit(page_no=None, section_path="body", text=(msg.body or "").strip()))
        for att in msg.attachments:
            name = getattr(att, "longFilename", None) or "attachment"
            units.append(ParsedUnit(page_no=None, section_path=f"attachment:{name}",
                                    text=f"[attachment: {name}]"))
    except Exception as exc:  # noqa: BLE001
        log.warning("email_parse_failed", error=str(exc))
    return ParsedDocument(units=units, page_count=1)


def _parse_text(ocr_pages: list[OCRPage], data: bytes) -> ParsedDocument:
    if ocr_pages:
        units = [ParsedUnit(page_no=p.page_no, section_path=None, text=p.text)
                 for p in ocr_pages if p.text.strip()]
        return ParsedDocument(units=units, page_count=len(ocr_pages))
    text = data.decode("utf-8", errors="replace").strip()
    paras = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    return ParsedDocument(units=[ParsedUnit(page_no=1, section_path=None, text=p) for p in paras],
                          page_count=1)


def _as_block(text: str):
    from app.core.ocr import OCRBlock

    return OCRBlock(text=text or "", bbox=[])
